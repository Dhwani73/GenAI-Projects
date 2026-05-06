# ingestion/parser.py

import fitz
from docx import Document
from docx.oxml.ns import qn
import os
from utils.logger import get_logger

logger = get_logger()


def extract_docx_hyperlinks(doc) -> dict:
    """
    Extracts all hyperlinks from a DOCX file.
    Returns a dict mapping relationship ID to URL.
    """
    hyperlinks = {}
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype:
            hyperlinks[rel.rId] = rel._target
    return hyperlinks


def extract_docx_tables(doc) -> str:
    """
    Extracts all tables from a DOCX file.
    Converts each table to readable pipe-separated text
    so the LLM can read and reason about it.
    """
    if not doc.tables:
        return ""

    table_text = "\n\n--- TABLES IN DOCUMENT ---\n"

    for table_idx, table in enumerate(doc.tables):
        try:
            table_text += f"\nTable {table_idx + 1}:\n"

            rows_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Clean cell text
                    cell_text = cell.text.strip().replace("\n", " ")
                    row_data.append(cell_text)
                rows_data.append(row_data)

            if not rows_data:
                continue

            # First row as header
            headers = rows_data[0]
            col_widths = [
                max(len(str(row[i])) if i < len(row) else 0
                    for row in rows_data)
                for i in range(len(headers))
            ]

            # Header row
            header_line = " | ".join(
                str(h).ljust(col_widths[i])
                for i, h in enumerate(headers)
            )
            separator = "-+-".join("-" * w for w in col_widths)
            table_text += header_line + "\n"
            table_text += separator + "\n"

            # Data rows
            for row in rows_data[1:]:
                row_line = " | ".join(
                    str(row[i]).ljust(col_widths[i])
                    if i < len(row) else " " * col_widths[i]
                    for i in range(len(headers))
                )
                table_text += row_line + "\n"

            table_text += "\n"

        except Exception as e:
            logger.warning(f"Could not extract table {table_idx + 1}: {e}")
            continue

    return table_text


def parse_pdf(file_path: str) -> str:
    logger.info(f"Parsing PDF: {file_path}")
    text = ""
    try:
        doc = fitz.open(file_path)
        logger.debug(f"PDF has {len(doc)} pages")

        for page_num, page in enumerate(doc):
            page_text = f"\n--- Page {page_num + 1} ---\n"

            # ── Extract text with inline links ──
            blocks = page.get_text("dict")["blocks"]
            links = page.get_links()

            # Map link rectangles to URLs
            link_map = {}
            for link in links:
                if "uri" in link:
                    link_map[fitz.Rect(link["from"])] = link["uri"]

            for block in blocks:
                if block["type"] != 0:
                    continue
                for line in block["lines"]:
                    line_text = ""
                    for span in line["spans"]:
                        span_text = span["text"].strip()
                        if not span_text:
                            continue

                        span_rect = fitz.Rect(span["bbox"])
                        matched_url = None

                        for link_rect, uri in link_map.items():
                            span_center_x = (span_rect.x0 + span_rect.x1) / 2
                            span_center_y = (span_rect.y0 + span_rect.y1) / 2
                            if link_rect.contains(
                                fitz.Point(span_center_x, span_center_y)
                            ):
                                matched_url = uri
                                break

                        if matched_url:
                            line_text += f"{span_text} ({matched_url}) "
                        else:
                            line_text += span_text + " "

                    if line_text.strip():
                        page_text += line_text.strip() + "\n"

            # ── Extract tables from page ────────
            try:
                tables = page.find_tables()
                if tables.tables:
                    page_text += "\n--- TABLES ON THIS PAGE ---\n"
                    for table_idx, table in enumerate(tables.tables):
                        page_text += f"\nTable {table_idx + 1}:\n"
                        extracted = table.extract()

                        if not extracted:
                            continue

                        # Header row
                        headers = [
                            str(cell).strip() if cell else ""
                            for cell in extracted[0]
                        ]
                        col_widths = [
                            max(
                                len(str(extracted[r][c]).strip() if extracted[r][c] else "")
                                for r in range(len(extracted))
                            )
                            for c in range(len(headers))
                        ]

                        header_line = " | ".join(
                            headers[i].ljust(col_widths[i])
                            for i in range(len(headers))
                        )
                        separator = "-+-".join(
                            "-" * w for w in col_widths
                        )
                        page_text += header_line + "\n"
                        page_text += separator + "\n"

                        # Data rows
                        for row in extracted[1:]:
                            row_line = " | ".join(
                                str(row[i]).strip().ljust(col_widths[i])
                                if i < len(row) and row[i]
                                else " " * col_widths[i]
                                for i in range(len(headers))
                            )
                            page_text += row_line + "\n"

                        page_text += "\n"

            except Exception as e:
                logger.warning(f"Table extraction failed on page {page_num + 1}: {e}")

            text += page_text

        doc.close()
        logger.info(f"PDF parsed — {len(text)} characters extracted")
    except Exception as e:
        logger.error(f"Failed to parse PDF: {e}")
        raise ValueError(f"Failed to parse PDF: {e}")
    return text.strip()


def parse_docx(file_path: str) -> str:
    logger.info(f"Parsing DOCX: {file_path}")
    text = ""
    try:
        doc = Document(file_path)
        hyperlinks = extract_docx_hyperlinks(doc)
        logger.debug(f"Found {len(hyperlinks)} hyperlinks in document")

        # ── Extract paragraphs ──────────────
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            style_name = para.style.name.lower()
            para_text = ""

            for child in para._element:
                if child.tag.endswith("}r"):
                    run_text = ""
                    for t in child.findall(f".//{qn('w:t')}"):
                        run_text += t.text or ""
                    para_text += run_text

                elif child.tag.endswith("}hyperlink"):
                    r_id = child.get(qn("r:id"))
                    link_text = ""
                    for t in child.findall(f".//{qn('w:t')}"):
                        link_text += t.text or ""

                    if r_id and r_id in hyperlinks:
                        url = hyperlinks[r_id]
                        para_text += f"{link_text} ({url})"
                    else:
                        para_text += link_text

            if not para_text.strip():
                continue

            if "heading 1" in style_name:
                text += f"\n\n# {para_text}\n"
            elif "heading 2" in style_name:
                text += f"\n\n## {para_text}\n"
            elif "heading 3" in style_name:
                text += f"\n\n### {para_text}\n"
            elif "heading 4" in style_name:
                text += f"\n\n#### {para_text}\n"
            elif "list" in style_name or para.style.name.startswith("List"):
                text += f"  - {para_text}\n"
            else:
                text += f"{para_text}\n"

        # ── Extract tables ──────────────────
        table_text = extract_docx_tables(doc)
        if table_text:
            text += table_text
            logger.info(f"Extracted {len(doc.tables)} table(s) from DOCX")

        logger.info(f"DOCX parsed — {len(text)} characters extracted")

    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
        raise ValueError(f"Failed to parse DOCX: {e}")
    return text.strip()


def parse_document(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"Document type detected: {ext.upper()}")

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    else:
        logger.error(f"Unsupported file type: {ext}")
        raise ValueError(f"Unsupported file type: {ext}")